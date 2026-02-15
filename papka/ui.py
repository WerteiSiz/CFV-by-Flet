import flet as ft
import math
import os
import webbrowser
from datetime import datetime
from docx import Document
from docx.shared import Pt  # Для работы с размерами шрифта
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Для выравнивания текста
from openpyxl import Workbook
from flet.core.alignment import Alignment
from openpyxl.styles import Alignment as OpenPyXLAlignment, Border, Side
import pandas as pd
from sklearn.ensemble import BaggingClassifier
from sklearn.tree import DecisionTreeClassifier
# Глобальные параметры
R3 = 6378.0  # Радиус Земли
H = 35810.0  # Высота орбиты


def ygol_mesta(gradys_d, gradys_sh, tochka):
    global R3, H, delta_y, cos_delta_y, cosinusFi, cosinustri, y0, c, e, drop, res_ygol
    R3 = 6378.0
    H = 35810.0
    delta_y = gradys_d - tochka
    cos_delta_y = math.cos(math.radians(delta_y))
    cosinusFi = math.cos(math.radians(gradys_sh))
    cosinustri = cosinusFi * cos_delta_y
    y0 = (R3 / (R3 + H))
    c = cosinustri - y0
    e = math.sqrt(1 - (cosinustri ** 2))
    drop = c / e
    res = math.atan(drop)
    res_ygol = math.degrees(res)
    return res_ygol

def azimut(gradys_d, gradys_sh, tochka):
    tangens_delta_y = math.tan(math.radians(delta_y))
    sinus_fi = math.sin(math.radians(gradys_sh))
    drop2 = tangens_delta_y / sinus_fi
    arct = math.atan(drop2)
    arct_res = math.degrees(arct) + 180
    return arct_res

def dalnost(gradys_d, gradys_sh, tochka):
    global d, sl_pdk, koren, drop3
    sl_pdk = 1 + y0 ** 2 - 2 * y0 * cosinustri
    koren = math.sqrt(sl_pdk)
    drop3 = koren / y0
    d = R3 * drop3
    return d

def poteri(chastota, diametr, k, tochka):
    global lenth
    global d, sl_pdk, koren, drop3
    sl_pdk = 1 + y0 ** 2 - 2 * y0 * cosinustri
    koren = math.sqrt(sl_pdk)
    drop3 = koren / y0
    d = R3 * drop3
    lenth = (3 * 10 ** 8 / (chastota * 10 ** 9))
    drop4 = math.log10((4 * math.pi * d * 1000) / lenth)
    L = 20 * drop4
    return L

def factorG(chastota, k, diametr):
    lenth = (3 * 10 ** 8 / (chastota * 10 ** 9))
    drop5 = ((math.pi * diametr) / lenth) ** 2
    G = 10 * math.log10(k * drop5)
    return G

def atmosphere(chastota, gradys_d, gradys_sh, tochka, T, p, P):
    global delta1, delta2, deltas, l1, l2, drooop, lh20, l0, drop16, drop11, drop14
    h0 = 6  # km in meters
    h3 = 0.02
    l1 = (5.98 / (math.sin(math.radians(ygol_mesta(gradys_d, gradys_sh, tochka)))))
    delta1 = (0.0126 / T ** 0.75)
    delta2 = (0.035 / T ** 0.75)
    deltas = 0.0153 * (1 + 0.0046 * p) / (T ** 0.5)
    drop6 = 1 / (chastota - 183.3) ** 2
    drop7 = 1 / (chastota - 323.8) ** 2
    drop8 = 3 / ((chastota - 22.3) ** 2 + 3)
    hh20 = drop6 + drop7 + 2.2 + drop8
    l2 = ((hh20 - h3) / (math.sin(math.radians(ygol_mesta(gradys_d, gradys_sh, tochka)))))
    drooop = (hh20 - h3)
    drop9 = (chastota ** 2 / T ** 2)
    drop10 = P * delta2
    drop11 = (2 - (chastota / 30)) ** 2
    drop12 = (drop10 / drop11 + (drop10 ** 2))
    drop13 = P * delta1
    drop14 = chastota ** 2 / 900
    drop15 = (drop13 / (drop13 ** 2 + drop14))
    drop16 = (2 + (chastota / 30)) ** 2
    drop17 = (drop10 / (drop16 + drop10 ** 2))
    l0 = 0.321 * P * drop9 * (drop12 + drop15 + drop17)
    drop18 = P * deltas / ((-0.741 + (chastota / 30)) ** 2 + (deltas * P) ** 2)
    drop19 = P * deltas / ((0.741 + (chastota / 30)) ** 2 + (deltas * P) ** 2)
    drop20 = 644 / T
    drop21 = ((5.72 * p * chastota ** 2 * e ** drop20) / T ** 2.5)
    drop22 = (P * deltas * p) / T
    lh20 = (drop18 + 0.0163 * chastota ** 2 * drop22 + drop19) * drop21
    latmosphere = ((l0 * l1) + (lh20 * l2))
    return latmosphere


COLORS = {
    "ACCENT": "#E2D4F0",  # Контейнер полей ввода и сверху сапр (лаванда)
    "dark_purple": "#9932CC",  # Тёмная орхидея - текст в полях
    "Fon": "#FFF9F0",  # Молочный
    "accent": "#A78BFA",  # Огранка всего
    "graphit": "#414A4C",
    "light_blue": "#1F91DC"
}


def main(page: ft.Page):

    page.title = "САПР VSAT"
    page.bgcolor = COLORS["Fon"]
    page.padding = 20
    page.scroll = ft.ScrollMode.AUTO  # Включение автоматической прокрутки
    page.auto_scroll = True  # Автоматическая прокрутка при добавлении новых элементов

    page.appbar = ft.AppBar(
        title=ft.Text("САПР VSAT", color=COLORS["dark_purple"]),
        bgcolor=COLORS["ACCENT"],
        center_title=True,
        toolbar_height=70
    )
    def show_result(e):
        try:
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)
            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)

            result_text = f"""
╔══════════════════════════════╗
║    РЕЗУЛЬТАТЫ ПОЗИЦИОНИРОВАНИЯ    ║
╠══════════════════════════════╣
║  Входные данные:             ║
║  • Широта: {gradys_sh:>12}°      ║
║  • Долгота: {gradys_d:>11}°      ║
║  • Точка стояния: {tochka:>6}°      ║
╠══════════════════════════════╣
║  Результаты:                 ║
║  • Угол места: {angle_of_elevation:>8.2f}°      ║
║  • Азимут: {azimuth_angle:>12.2f}°      ║
║  • Дальность: {distance:>10.2f} км    ║
╚══════════════════════════════╝
"""

            result_card = ft.Card(
                content=ft.Container(
                    content=ft.Column([
                        ft.Text("Результаты расчёта",
                                size=18,
                                weight="bold",
                                color=COLORS["accent"]),
                        ft.Divider(height=1, color=COLORS["ACCENT"]),
                        ft.Text(result_text,
                                font_family="Comic Sans",
                                size=14,
                                color=COLORS["dark_purple"])
                    ], spacing=10),
                    padding=20,
                    width=400
                ),
                elevation=8,
                color=COLORS["ACCENT"],
                margin=10
            )

            result_container.controls.clear()
            result_container.controls.append(
                ft.Row([result_card], alignment=ft.MainAxisAlignment.CENTER)
            )
            page.update()

        except ValueError:
            show_save_message("Ошибка: проверьте введенные данные", is_error=True)
            page.update()
    def show_result_e(e):
        try:
            # Получаем и проверяем входные данные
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)
            chastota = float(chastota_input.value)
            diametr = float(diametr_input.value)
            k = float(k_input.value)
            P = float(P_input.value)
            T = float(T_input.value)
            p = float(p_input.value)

            # Вычисляем параметры
            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)
            signal_loss = round(poteri(chastota, diametr, k, tochka), 3)
            gain = round(factorG(chastota, k, diametr), 3)
            atmos = round(atmosphere(chastota, gradys_d, gradys_sh, tochka, T, p, P), 3)
            lenth = 0.3 / chastota  # Пример вычисления длины волны

            # Форматируем результат
            result_text = f"""
    ╔════════════════════════════════════════════╗
    ║           РЕЗУЛЬТАТЫ РАСЧЕТОВ             ║
    ╚════════════════════════════════════════════╝

    ▌ Входные параметры:
    ├────────────────────────────────────────────
    │ ▪ Широта: {gradys_sh}°
    │ ▪ Долгота: {gradys_d}°
    │ ▪ Точка стояния: {tochka}°
    │ ▪ Частота: {chastota} м
    │ ▪ Диаметр антенны: {diametr} м
    │ ▪ Коэф. использования: {k}
    │ ▪ Давление: {P} кПа
    │ ▪ Температура: {T} °K
    │ ▪ Влажность: {p} г/м³

    ▌ Позиционирование:
    ├────────────────────────────────────────────
    │ ▪ Угол места: {angle_of_elevation:.2f}°
    │ ▪ Азимут: {azimuth_angle:.2f}°
    │ ▪ Наклонная дальность: {distance:.2f} км
    │ ──────────────────────────

    ▌ Энергетические параметры:
    ├────────────────────────────────────────────
    │ ▪ Длина волны: {lenth:.4f} м
    │ ▪ Потери в свободном пространстве: {signal_loss:.2f} дБ
    │ ▪ Потери в атмосфере: {atmos:.2f} дБ
    │ ▪ Коэффициент усиления: {gain:.2f} дБ
    ╰────────────────────────────────────────────
    """

            # Создаем карточку с результатами
            result_card = ft.Card(
                content=ft.Container(
                    content=ft.Column([
                        ft.Text("Результаты расчета",
                                size=20,
                                weight=ft.FontWeight.BOLD,
                                color=COLORS["accent"]),
                        ft.Divider(height=10, color="transparent"),
                        ft.Text(result_text,
                                font_family="Courier New",
                                size=14,
                                color=COLORS["dark_purple"],
                                selectable=True)
                    ], spacing=5),
                    padding=25,
                    width=450
                ),
                elevation=10,
                color=COLORS["ACCENT"],
                margin=15,
                shadow_color=COLORS["dark_purple"]
            )

            # Очищаем и обновляем контейнер с результатами
            result_container.controls.clear()
            result_container.controls.append(
                ft.Row([result_card], alignment=ft.MainAxisAlignment.CENTER)
            )
            page.update()

        except ValueError as ve:
            show_save_message(f"Ошибка ввода: {str(ve)}", is_error=True)
            page.update()
        except Exception as ex:
            show_save_message(f"Ошибка расчета: {str(ex)}", is_error=True)
            page.update()
    def show_result_e1(e):
        try:
            # Получаем и проверяем входные данные
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)
            chastota1 = float(chastota_input.value)
            diametr = float(diametr_input.value)
            k = float(k_input.value)
            P = float(P_input.value)
            T = float(T_input.value)
            p = float(p_input.value)

            # Вычисляем параметры
            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)
            signal_loss = round(poteri(chastota1, diametr, k, tochka), 3)
            gain = round(factorG(chastota1, k, diametr), 3)
            atmos = round(atmosphere(chastota1, gradys_d, gradys_sh, tochka, T, p, P), 3)
            lenth = 0.3 / chastota1  # Пример вычисления длины волны

            # Форматируем результат
            result_text = f"""
    ╔════════════════════════════════════════════╗
    ║           РЕЗУЛЬТАТЫ РАСЧЕТОВ             ║
    ╚════════════════════════════════════════════╝

    ▌ Входные параметры:
    ├────────────────────────────────────────────
    │ ▪ Широта: {gradys_sh}°
    │ ▪ Долгота: {gradys_d}°
    │ ▪ Точка стояния спутника: {tochka}°
    │ ▪ Диаметр антенны: {diametr} м
    │ ▪ Частота: {chastota1} м
    │ ▪ Коэф. использования: {k}
    │ ▪ Давление: {P} кПа
    │ ▪ Температура: {T} °K
    │ ▪ Влажность: {p} г/м³

    ▌ Позиционирование:
    ├────────────────────────────────────────────
    │ ▪ Угол места: {angle_of_elevation:.2f}°
    │ ▪ Азимут: {azimuth_angle:.2f}°
    │ ▪ Наклонная дальность: {distance:.2f} км
    │ ──────────────────────────

    ▌ Энергетические параметры:
    ├────────────────────────────────────────────
    │ ▪ Длина волны: {lenth:.4f} м
    │ ▪ Потери в свободном пространстве: {signal_loss:.2f} дБ
    │ ▪ Потери в атмосфере: {atmos:.2f} дБ
    │ ▪ Коэффициент усиления: {gain:.2f} дБ
    ╰────────────────────────────────────────────
    """

            # Создаем карточку с результатами
            result_card = ft.Card(
                content=ft.Container(
                    content=ft.Column([
                        ft.Text("Результаты расчета",
                                size=20,
                                weight=ft.FontWeight.BOLD,
                                color=COLORS["accent"]),
                        ft.Divider(height=10, color="transparent"),
                        ft.Text(result_text,
                                font_family="Courier New",
                                size=14,
                                color=COLORS["dark_purple"],
                                selectable=True)
                    ], spacing=5),
                    padding=25,
                    width=450
                ),
                elevation=10,
                color=COLORS["ACCENT"],
                margin=15,
                shadow_color=COLORS["dark_purple"]
            )

            # Очищаем и обновляем контейнер с результатами
            result_container.controls.clear()
            result_container.controls.append(
                ft.Row([result_card], alignment=ft.MainAxisAlignment.CENTER)
            )
            page.update()

        except ValueError as ve:
            show_save_message(f"Ошибка ввода: {str(ve)}", is_error=True)
            page.update()
        except Exception as ex:
            show_save_message(f"Ошибка расчета: {str(ex)}", is_error=True)
            page.update()
    def load_model():
        # Загрузка набора данных
        data = pd.read_excel(r'C:\TIP\projects\CFV\save_document\data_excel.xlsx')

        # Определение целевых и признаковых столбцов
        columns_target = 'Диаметр'
        columns_train = ['Тип орбитальной группировки', 'Частота', 'С\Ш']

        # Подготовка данных
        X = data[columns_train].copy()
        Y = data[columns_target]

        # Кодирование столбца
        X['Тип орбитальной группировки'] = X['Тип орбитальной группировки'].map({'GEO': 0})

        # Инициализация и обучение модели
        base_classifier = DecisionTreeClassifier()
        bagging_classifier = BaggingClassifier(estimator=base_classifier, n_estimators=50, random_state=42)
        bagging_classifier.fit(X, Y)

        return bagging_classifier, columns_train

        # Загружаем модель при старте приложения

    model, feature_columns = load_model()

    def antenna_diameter_decision(e):
        """СППР - расчет диаметра антенны"""

        # Функция для обработки нажатия кнопки "Рассчитать"
        def calculate_click(e):
            try:
                # Получаем данные из полей ввода
                freq = int(frequency_input.value)
                orbit = orbit_dropdown.value  # Получаем выбранное значение из dropdown
                s_h = float(sn_ratio_input.value)

                # Создаем DataFrame с введенными данными
                user_data = pd.DataFrame([[orbit, freq, s_h]], columns=feature_columns)
                user_data['Тип орбитальной группировки'] = user_data['Тип орбитальной группировки'].map(
                    {'GEO': 0, 'MEO': 1, 'LEO': 2}  # Добавляем маппинг для всех типов орбит
                )

                # Делаем предсказание
                prediction = model.predict(user_data)

                # Определяем диаметр рефлектора
                if prediction[0] == 1:
                    diameter = "Диаметр рефлектора: 0.98 м"
                elif prediction[0] == 2:
                    diameter = "Диаметр рефлектора: 1.2 м"
                elif prediction[0] == 3:
                    diameter = "Диаметр рефлектора: 1.8 м"
                else:
                    diameter = "Неизвестный диаметр"

                # Показываем результат
                show_save_message(diameter, is_error=False)
                page.update()

            except Exception as ex:
                show_save_message(f"Ошибка: {str(ex)}", is_error=True)
                page.update()

        # Очищаем страницу и добавляем меню
        page.controls.clear()
        page.add(menubar)

        # Настройка AppBar
        page.appbar = ft.AppBar(
            title=ft.Text("Система поддержки принятия решений",
                          color=COLORS["dark_purple"]),
            center_title=True,
            bgcolor=COLORS["ACCENT"],
            toolbar_height=70
        )

        # Создаем элементы интерфейса
        title = ft.Container(
            ft.Text(
                value="Предсказание диаметра рефлектора",
                text_align="center",
                size=24,
                color=COLORS["dark_purple"],
                weight=ft.FontWeight.BOLD
            ),
            alignment=ft.alignment.center,
            margin=ft.margin.only(bottom=20)
        )

        # Общий стиль для полей ввода
        common_input_style = {
            "border_color": COLORS["accent"],
            "border_radius": 10,
            "border_width": 1.5,
            "bgcolor": ft.Colors.WHITE,
            "label_style": ft.TextStyle(color=COLORS["dark_purple"]),
            "color": COLORS["dark_purple"],
            "width": 300
        }

        frequency_input = ft.TextField(
            label="Частота (ГГц)",
            **common_input_style
        )

        # Выпадающий список для выбора типа орбиты
        orbit_dropdown = ft.Dropdown(
            label="Тип орбиты",
            options=[
                ft.dropdown.Option("GEO"),
                ft.dropdown.Option("MEO"),
                ft.dropdown.Option("LEO"),
            ],
            value="GEO",  # Значение по умолчанию
            **common_input_style
        )

        sn_ratio_input = ft.TextField(
            label="Соотношение С/Ш",
            **common_input_style
        )

        # Стиль для кнопок
        btn_style = ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=12),
            padding=18,
            side=ft.BorderSide(1, COLORS["accent"]),
            bgcolor=ft.Colors.WHITE,
            overlay_color=ft.Colors.TRANSPARENT
        )

        # Кнопка расчета
        calculate_btn = ft.ElevatedButton(
            text="Рассчитать",
            icon=ft.Icons.CALCULATE,
            on_click=calculate_click,
            style=btn_style,
            width=200
        )

        # Карточка с формой ввода
        input_card = ft.Card(
            content=ft.Container(
                content=ft.Column(
                    [
                        title,
                        frequency_input,
                        orbit_dropdown,  # Используем dropdown вместо TextField
                        sn_ratio_input,
                        ft.Divider(height=20, color="transparent"),
                        ft.Row([calculate_btn], alignment=ft.MainAxisAlignment.CENTER)
                    ],
                    alignment=ft.MainAxisAlignment.CENTER,
                    horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                    spacing=20
                ),
                padding=30,
                width=400
            ),
            elevation=8,
            color=COLORS["ACCENT"],
            margin=10
        )

        # Добавляем элементы на страницу
        page.add(
            ft.Column(
                [
                    ft.Row([input_card], alignment=ft.MainAxisAlignment.CENTER),
                    save_message_container
                ],
                spacing=20,
                expand=True,
            )
        )
        page.update()
    def show_save_message(message, is_error=False):
        """Показывает красивое сообщение о сохранении"""
        save_message.value = message
        save_message.color = ft.Colors.PINK_400 if is_error else COLORS["light_blue"]
        save_message_container.visible = True
        page.update()

        # Автоматическое скрытие сообщения через 3 секунды
        def hide_message():
            save_message_container.visible = False
            page.update()

        import threading
        timer = threading.Timer(3.0, hide_message)
        timer.start()

    # Функции сохранения
    def save_result_txt(e):
        try:
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)
            chastota = float(chastota_input.value)
            diametr = float(diametr_input.value)
            k = float(k_input.value)
            P = float(P_input.value)
            T = float(T_input.value)
            p = float(p_input.value)

            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)
            signal_loss = round(poteri(chastota, diametr, k, tochka), 3)
            gain = round(factorG(chastota, k, diametr), 3)
            atmos = round(atmosphere(chastota, gradys_d, gradys_sh, tochka, T, p, P), 3)

            result_text = f"""
    ╔════════════════════════════════╗
    ║       ОТЧЕТ О РАСЧЕТАХ        ║
    ╠════════════════════════════════╣
    ║  ВХОДНЫЕ ДАННЫЕ               ║
    ║  • Широта: {gradys_sh:>12}°       ║
    ║  • Долгота: {gradys_d:>11}°       ║
    ║  • Точка стояния: {tochka:>6}°       ║
    ║  • Диаметр антенны: {diametr:>5.2f} м     ║
    ║  • Коэффициент: {k:>10.2f}       ║
    ║  • Давление: {P:>11.2f} кПа    ║
    ║  • Температура: {T:>8.2f} °K   ║
    ║  • Влажность: {p:>10.2f} г/м³  ║
    ╠════════════════════════════════╣
    ║  РЕЗУЛЬТАТЫ                   ║
    ║  • Угол места: {angle_of_elevation:>8.2f}°       ║
    ║  • Азимут: {azimuth_angle:>12.2f}°       ║
    ║  • Дальность: {distance:>10.2f} км     ║
    ║  • Потери сигнала: {signal_loss:>6.2f} дБ ║
    ║  • Потери в атмосфере: {atmos:>4.2f} дБ ║
    ║  • Коэффициент усиления: {gain:>2.2f} дБ ║
    ╚════════════════════════════════╝
    """

            with open("Результаты_Расчётов.txt", "w", encoding="utf-8") as file:
                file.write(result_text)

            show_save_message("Данные сохранены в Результаты_Расчётов.txt")

        except ValueError:
            show_save_message("Ошибка: проверьте корректность данных", is_error=True)

    def save_result_to_word(e):
        try:
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)

            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)

            result_text = f"""Результаты Вычислений:
                   Ваши входные данные:
                       Градусы широты: {gradys_sh}°
                       Градусы Долготы: {gradys_d}°
                       Значение Точки стояния спутника: {tochka}°


                   Результаты Позиционирования:
                       Угол места: {angle_of_elevation:.2f}°
                       Азимут: {azimuth_angle:.2f}°
                       Наклонная дальность: {distance:.2f} км
                           ~~~~~~~~~~~~~~~~~~
                Спасибо, что пользуетесь нашей системой расчётов!"""

            user_input = result_text

            desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            file_path = os.path.join(desktop_path, "output.docx")

            doc = Document()
            doc.add_paragraph(user_input)
            doc.save(file_path)

            with open("Результаты_Расчётов.doc", "w") as file:
                file.write(result_text)
            show_save_message("Данные сохранены в файл Результаты_Расчётов.doc")
        except ValueError:
            show_save_message("Пожалуйста, введите корректные данные", is_error=True)

    def save_result_of_positioning_txt(e):
        """Сохранение результатов позиционирования в TXT файл"""
        try:
            # Получаем данные из полей ввода
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)

            # Выполняем расчеты
            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)

            # Форматируем результат с рамкой и разделителями
            result_text = f"""
    ╔══════════════════════════════════════════════╗
    ║           РЕЗУЛЬТАТЫ ПОЗИЦИОНИРОВАНИЯ        ║
    ╠══════════════════════════════════════════════╣
    ║  ВХОДНЫЕ ДАННЫЕ:                            ║
    ║  • Широта: {gradys_sh:>12}°                 ║
    ║  • Долгота: {gradys_d:>11}°                 ║
    ║  • Точка стояния спутника: {tochka:>6}°     ║
    ╠══════════════════════════════════════════════╣
    ║  РЕЗУЛЬТАТЫ:                                ║
    ║  • Угол места: {angle_of_elevation:>8.2f}°  ║
    ║  • Азимут: {azimuth_angle:>12.2f}°          ║
    ║  • Наклонная дальность: {distance:>10.2f} км║
    ╠══════════════════════════════════════════════╣
    ║  Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')} ║
    ╚══════════════════════════════════════════════╝
    """

            # Сохраняем в файл с кодировкой UTF-8
            with open("Результаты_Позиционирования.txt", "w", encoding="utf-8") as file:
                file.write(result_text)

            # Показываем уведомление об успешном сохранении
            show_save_message("Данные сохранены в Результаты_Позиционирования.txt")

        except ValueError:
            # Показываем сообщение об ошибке
            show_save_message("Ошибка: проверьте корректность введенных данных", is_error=True)
        except Exception as ex:
            # Обработка других возможных ошибок
            show_save_message(f"Ошибка при сохранении: {str(ex)}", is_error=True)

    def save_result_of_positioning_to_word(e):
        """Сохранение результатов позиционирования в документ Word"""
        try:
            # Получаем и проверяем данные
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)

            # Выполняем расчеты
            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)

            # Создаем документ Word
            doc = Document()

            # Добавляем заголовок с форматированием
            title = doc.add_paragraph()
            title_run = title.add_run("ОТЧЕТ О РАСЧЕТАХ ПОЗИЦИОНИРОВАНИЯ")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Добавляем дату
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_run.italic = True
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Раздел входных данных
            doc.add_paragraph("Входные данные:", style='Heading 2')
            input_data = [
                ("Широта", f"{gradys_sh}°"),
                ("Долгота", f"{gradys_d}°"),
                ("Точка стояния спутника", f"{tochka}°")
            ]
            for name, value in input_data:
                doc.add_paragraph(f"{name}: {value}", style='List Bullet')

            # Раздел результатов
            doc.add_paragraph("Результаты позиционирования:", style='Heading 2')
            results = [
                ("Угол места", f"{angle_of_elevation:.2f}°"),
                ("Азимут", f"{azimuth_angle:.2f}°"),
                ("Наклонная дальность", f"{distance:.2f} км")
            ]
            for name, value in results:
                doc.add_paragraph(f"{name}: {value}", style='List Bullet')

            # Добавляем подпись
            doc.add_paragraph()
            sign_para = doc.add_paragraph("Спасибо за использование нашей системы расчетов!")
            sign_para.italic = True
            sign_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Сохраняем документ
            desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
            file_path = os.path.join(desktop_path, "Результаты_позиционирования.docx")
            doc.save(file_path)

            # Уведомление об успешном сохранении
            show_save_message(f"Файл Результаты_позиционирования.docx сохранен!")

        except ValueError:
            show_save_message("Ошибка: проверьте корректность введенных данных", is_error=True)
        except Exception as ex:
            show_save_message(f"Ошибка при сохранении в Word: {str(ex)}", is_error=True)
    def save_result_excel(e):
        try:
            gradys_d = float(gradys_d_input.value)
            gradys_sh = float(gradys_sh_input.value)
            tochka = float(tochka_input.value)
            chastota = float(chastota_input.value)
            chastota1 = float(chastota_input.value)
            diametr = float(diametr_input.value)
            k = float(k_input.value)
            P = float(P_input.value)
            T = float(T_input.value)
            p = float(p_input.value)

            angle_of_elevation = ygol_mesta(gradys_d, gradys_sh, tochka)
            azimuth_angle = azimut(gradys_d, gradys_sh, tochka)
            distance = dalnost(gradys_d, gradys_sh, tochka)
            signal_loss = round(poteri(chastota, diametr, k, tochka), 3)
            signal_loss1 = round(poteri(chastota1, diametr, k, tochka), 3)
            signal_loss2 = round(poteri(chastota1, diametr, k, tochka), 3)
            signal_loss3 = round(poteri(chastota1, diametr, k, tochka), 3)
            gain = round(factorG(chastota, k, diametr), 3)
            atmos = round(atmosphere(chastota, gradys_d, gradys_sh, tochka, T, p, P), 3)

            # Создаем новую книгу Excel
            wb = Workbook()

            # Выбираем активный лист
            ws = wb.active

            # Устанавливаем заголовки столбцов
            ws['A1'] = 'Параметр/Антенна'
            ws['B1'] = 'General Dynamics 0,98 м'
            ws['C1'] = 'General Dynamics 1,2 м'
            ws['D1'] = 'General Skyware 1,8 '

            # Устанавливаем значения в столбце A
            ws['A2'] = 'Тип линии'
            ws['A3'] = 'Рабочая частота, ГГц'
            ws['A4'] = 'Наклонная дальность, км'
            ws['A5'] = 'Диаметр антенны'
            ws['A6'] = 'Потери в свободном пространстве, Дб'

            # Устанавливаем значения в столбце B
            ws['B2'] = f'Линия «вверх» ЗС-ИЗС \nЛиния «вниз» ИЗС-ЗС '  # OK
            ws['B2'].alignment = Alignment(wrapText=True)
            ws['B3'] = f'{chastota:.2f}\n{chastota1:.2f}'  # OK
            ws['B3'].alignment = Alignment(wrapText=True)
            ws['B4'] = f'{distance:.2f}'  # OK
            ws['B5'] = f'{diametr:.2f} м'  # OK
            ws['B6'] = f'{signal_loss:.2f}\n{signal_loss1:.2f}'  # OK
            ws['B6'].alignment = Alignment(wrapText=True)

            # Устанавливаем значения в столбце C
            ws['C2'] = f'Линия «вверх» ЗС-ИЗС \nЛиния «вниз» ИЗС-ЗС '  # OK
            ws['C2'].alignment = Alignment(wrapText=True)
            ws['C3'] = f'{chastota:.2f}\n{chastota1:.2f}'
            ws['C3'].alignment = Alignment(wrapText=True)
            ws.merge_cells('B4:D4')
            ws['B4'] = f'{distance:.2f}'
            ws['C5'] = f'{diametr:.2f} м'
            ws['C6'] = f'{signal_loss:.2f}\n{signal_loss1:.2f}'
            ws['C6'].alignment = Alignment(wrapText=True)

            # Устанавливаем значения в столбце D
            ws['D2'] = f'Линия «вверх» ЗС-ИЗС \nЛиния «вниз» ИЗС-ЗС '  # OK
            ws['D2'].alignment = Alignment(wrapText=True)
            ws['D3'] = f'{chastota:.2f}\n{chastota1:.2f}'
            ws['D3'].alignment = Alignment(wrapText=True)
            ws['D5'] = f'{diametr:.2f} м'
            ws['D6'] = f'{signal_loss:.2f}\n{signal_loss1:.2f}'
            ws['D6'].alignment = Alignment(wrapText=True)

            # Сохраняем книгу Excel
            wb.save('Результаты_Расчётов.xlsx')
            show_save_message("Данные сохранены в файл Результаты_Расчётов.xlsx")
        except ValueError:
            show_save_message("Пожалуйста, введите корректные данные", is_error=True)

    def energy_calculation_up(e):
        """Энергетический расчет линии ЗС-ИЗС"""
        page.controls.clear()

        # Настройка AppBar
        page.appbar = ft.AppBar(
            title=ft.Text("Энергетический расчёт (ЗС-ИЗС)",
                          color=COLORS["dark_purple"]),
            center_title=True,
            bgcolor=COLORS["ACCENT"],
            toolbar_height=70
        )

        # Стиль для кнопок
        btn_style = ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=12),
            padding=18,
            side=ft.BorderSide(1, COLORS["accent"]),
            bgcolor=ft.Colors.WHITE,
            overlay_color=ft.Colors.TRANSPARENT
        )

        # Создаем карточку с полями ввода
        input_card = ft.Card(
            content=ft.Container(
                content=ft.ResponsiveRow(
                    [
                        ft.Container(gradys_d_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(gradys_sh_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(tochka_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(chastota_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(diametr_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(k_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(P_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(T_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(p_input, padding=15, col={"sm": 12, "md": 4}),
                    ],
                    spacing=20,
                ),
                padding=20,
            ),
            elevation=5,
            color=COLORS["ACCENT"],
            margin=ft.margin.only(bottom=30),
        )

        # Кнопки действий
        buttons_row = ft.Row(
            [
                ft.ElevatedButton(
                    text="Произвести расчёты",
                    icon=ft.Icons.CALCULATE,
                    style=btn_style,
                    width=200,
                    on_click=show_result_e
                ),
                ft.Container(
                    content=ft.PopupMenuButton(
                        items=[
                            ft.PopupMenuItem(
                                text="Сохранить в TXT",
                                icon=ft.Icons.TEXT_SNIPPET,
                                on_click=save_result_txt
                            ),
                            ft.PopupMenuItem(
                                text="Сохранить в Word",
                                icon=ft.Icons.DESCRIPTION,
                                on_click=save_result_to_word
                            ),
                            ft.PopupMenuItem(
                                text="Сохранить в Excel",
                                icon=ft.Icons.TABLE_CHART,
                                on_click=save_result_excel
                            ),
                        ],
                        icon=ft.Icons.SAVE,
                        tooltip="Выберите формат сохранения",
                        style=ft.ButtonStyle(
                            shape=ft.RoundedRectangleBorder(radius=12),
                            side=ft.BorderSide(1, COLORS["accent"]),
                            bgcolor=ft.Colors.WHITE,
                        )
                    ),
                    padding=10,
                    border_radius=12,
                    border=ft.border.all(1, COLORS["accent"]),
                    bgcolor=ft.Colors.WHITE,
                )
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=30
        )

        # Добавляем элементы на страницу
        page.add(
            ft.Column(
                [
                    menubar,
                    input_card,
                    buttons_row,
                    save_message_container,
                    result_container
                ],
                spacing=20,
                expand=True,
            )
        )
        page.update()

    def energy_calculation_down(e):
        """Энергетический расчет линии ИЗС-ЗС"""
        page.controls.clear()

        # Настройка AppBar
        page.appbar = ft.AppBar(
            title=ft.Text("Энергетический расчёт (ИЗС-ЗС)",
                          color=COLORS["dark_purple"]),
            center_title=True,
            bgcolor=COLORS["ACCENT"],
            toolbar_height=70
        )

        # Стиль для кнопок
        btn_style = ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=12),
            padding=18,
            side=ft.BorderSide(1, COLORS["accent"]),
            bgcolor=ft.Colors.WHITE,
            overlay_color=ft.Colors.TRANSPARENT
        )

        # Создаем карточку с полями ввода
        input_card = ft.Card(
            content=ft.Container(
                content=ft.ResponsiveRow(
                    [
                        ft.Container(gradys_d_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(gradys_sh_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(tochka_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(chastota1_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(diametr_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(k_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(P_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(T_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(p_input, padding=15, col={"sm": 12, "md": 4}),
                    ],
                    spacing=20,
                ),
                padding=20,
            ),
            elevation=5,
            color=COLORS["ACCENT"],
            margin=ft.margin.only(bottom=30),
        )

        # Кнопки действий
        buttons_row = ft.Row(
            [
                ft.ElevatedButton(
                    text="Произвести расчёты",
                    icon=ft.Icons.CALCULATE,
                    style=btn_style,
                    width=200,
                    on_click=show_result_e1
                ),
                ft.Container(
                    content=ft.PopupMenuButton(
                        items=[
                            ft.PopupMenuItem(
                                text="Сохранить в TXT",
                                icon=ft.Icons.TEXT_SNIPPET,
                                on_click=save_result_txt
                            ),
                            ft.PopupMenuItem(
                                text="Сохранить в Word",
                                icon=ft.Icons.DESCRIPTION,
                                on_click=save_result_to_word
                            ),
                            ft.PopupMenuItem(
                                text="Сохранить в Excel",
                                icon=ft.Icons.TABLE_CHART,
                                on_click=save_result_excel
                            ),
                        ],
                        icon=ft.Icons.SAVE,
                        tooltip="Выберите формат сохранения",
                        style=ft.ButtonStyle(
                            shape=ft.RoundedRectangleBorder(radius=12),
                            side=ft.BorderSide(1, COLORS["accent"]),
                            bgcolor=ft.Colors.WHITE,
                        )
                    ),
                    padding=10,
                    border_radius=12,
                    border=ft.border.all(1, COLORS["accent"]),
                    bgcolor=ft.Colors.WHITE,
                )
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=30
        )

        # Добавляем элементы на страницу
        page.add(
            ft.Column(
                [
                    menubar,
                    input_card,
                    buttons_row,
                    save_message_container,
                    result_container
                ],
                spacing=20,
                expand=True,
            )
        )
        page.update()

    def positioning(e):
        """Функция построения страницы позиционирования"""
        # Очищаем страницу и создаем меню
        page.controls.clear()


        # Настройка AppBar
        page.appbar = ft.AppBar(
            title=ft.Text("Позиционирование в пространстве",
                          color=COLORS["dark_purple"]),
            center_title=True,
            bgcolor=COLORS["ACCENT"],
            toolbar_height=70
        )

        # Стиль для кнопок
        btn_style = ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=12),
            padding=18,
            side=ft.BorderSide(1, COLORS["accent"]),
            bgcolor=ft.Colors.WHITE,
            overlay_color=ft.Colors.TRANSPARENT
        )

        # Создаем карточку с полями ввода
        input_card = ft.Card(
            content=ft.Container(
                content=ft.ResponsiveRow(
                    [
                        ft.Container(gradys_d_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(gradys_sh_input, padding=15, col={"sm": 12, "md": 4}),
                        ft.Container(tochka_input, padding=15, col={"sm": 12, "md": 4}),
                    ],
                    spacing=20,
                    run_spacing={"xs": 10, "sm": 15, "md": 20},
                ),
                padding=20,
            ),
            elevation=8,
            color=COLORS["ACCENT"],
            margin=ft.margin.only(bottom=30),
        )

        # Кнопки действий
        buttons_row = ft.Row(
            [
                ft.ElevatedButton(
                    text="Произвести расчёты",
                    icon=ft.Icons.CALCULATE,
                    style=btn_style,
                    width=200,
                    on_click=show_result
                ),
                ft.Container(
                    content=ft.PopupMenuButton(
                        items=[
                            ft.PopupMenuItem(
                                text="Сохранить в TXT",
                                icon=ft.Icons.TEXT_SNIPPET,
                                on_click=save_result_of_positioning_txt
                            ),
                            ft.PopupMenuItem(
                                text="Сохранить в Word",
                                icon=ft.Icons.DESCRIPTION,
                                on_click=save_result_of_positioning_to_word
                            ),

                        ],
                        icon=ft.Icons.SAVE,
                        tooltip="Выберите формат сохранения",
                        style=ft.ButtonStyle(
                            shape=ft.RoundedRectangleBorder(radius=12),
                            side=ft.BorderSide(1, COLORS["accent"]),
                            bgcolor=ft.Colors.WHITE,
                        )
                    ),
                    padding=10,
                    border_radius=12,
                    border=ft.border.all(1, COLORS["accent"]),
                    bgcolor=ft.Colors.WHITE,
                )
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            spacing=30
        )

        # Добавляем элементы на страницу
        page.add(
            ft.Column(
                [
                    menubar,
                    input_card,
                    buttons_row,
                    save_message_container,
                    result_container
                ],
                spacing=20,
                expand=True,
            )
        )
        page.update()
    # Создание полей ввода с темно-фиолетовым текстом
    def create_input_field(label):
        return ft.TextField(
            label=label,
            border_color=COLORS["accent"],
            border_radius=10,
            border_width=1.5,
            bgcolor=ft.Colors.WHITE,
            label_style=ft.TextStyle(color=COLORS["dark_purple"]),
            color=COLORS["dark_purple"],  # Цвет вводимого текста
            width=300
        )

    # Инициализация полей ввода
    gradys_d_input = create_input_field("Долгота (°)")
    gradys_sh_input = create_input_field("Широта (°)")
    tochka_input = create_input_field("Точка стояния (°)")
    chastota_input = create_input_field("Частота (ГГц)")
    chastota1_input = create_input_field("Частота (ГГц)")
    diametr_input = create_input_field("Диаметр антенны (м)")
    k_input = create_input_field("Коэффициент использования")
    P_input = create_input_field("Давление (кПа)")
    T_input = create_input_field("Температура (°K)")
    p_input = create_input_field("Влажность (г/м³)")

    # Контейнер для сообщений о сохранении
    save_message = ft.Text(color=COLORS["accent"], size=14)
    save_message_container = ft.Container(
        content=ft.Row(
            [
                save_message,
                ft.IconButton(
                    icon=ft.Icons.CLOSE,
                    on_click=lambda e: setattr(save_message_container, "visible", False) or page.update(),
                    icon_size=18,
                    tooltip="Закрыть"
                )
            ],
            alignment=ft.MainAxisAlignment.SPACE_BETWEEN
        ),
        padding=15,
        border_radius=10,
        bgcolor=COLORS["ACCENT"],
        visible=False,
        animate_opacity=300,
        width=400,
        margin=ft.margin.only(bottom=20)
    )

    # Стиль кнопок
    btn_style = ft.ButtonStyle(
        shape=ft.RoundedRectangleBorder(radius=12),
        padding=18,
        side=ft.BorderSide(1, COLORS["accent"]),
        bgcolor=ft.Colors.WHITE,
        overlay_color=ft.Colors.TRANSPARENT
    )

    # Выпадающее меню для сохранения с иконками
    save_menu = ft.PopupMenuButton(
        items=[
            ft.PopupMenuItem(
                text="Сохранить в TXT",
                icon=ft.Icons.TEXT_SNIPPET,
                on_click=save_result_txt
            ),
            ft.PopupMenuItem(
                text="Сохранить в Word",
                icon=ft.Icons.DESCRIPTION,
                on_click=save_result_to_word
            ),
            ft.PopupMenuItem(
                text="Сохранить в Excel",
                icon=ft.Icons.TABLE_CHART,
                on_click=save_result_excel
            ),
        ],
        icon=ft.Icons.SAVE,
        tooltip="Выберите формат сохранения",
        # Стиль для меню
        style=ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=12),
            side=ft.BorderSide(1, COLORS["accent"]),
            bgcolor=ft.Colors.WHITE,
        )
    )

    # Кнопки в строку с выпадающим меню сохранения
    buttons_row = ft.Row(
        [
            ft.ElevatedButton(
                text="Рассчитать",
                icon=ft.Icons.CALCULATE,
                style=btn_style,
                width=200,
                on_click=show_result
            ),
            ft.Container(
                content=save_menu,
                padding=10,
                border_radius=12,
                border=ft.border.all(1, COLORS["accent"]),
                bgcolor=ft.Colors.WHITE,
                on_hover=lambda e: setattr(e.control, "bgcolor", COLORS["ACCENT"]) or e.control.update()
            )
        ],
        alignment=ft.MainAxisAlignment.CENTER,
        spacing=30
    )

    # Контейнер для результатов
    result_container = ft.Column([], spacing=10)

    # Карточка с полями ввода
    input_card = ft.Card(
        content=ft.Container(
            content=ft.ResponsiveRow(
                [
                    ft.Container(gradys_d_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(gradys_sh_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(tochka_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(chastota_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(diametr_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(k_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(P_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(T_input, padding=15, col={"sm": 12, "md": 4}),
                    ft.Container(p_input, padding=15, col={"sm": 12, "md": 4}),
                ],
                spacing=20,
            ),
            padding=20,
        ),
        elevation=5,
        color=COLORS["ACCENT"],
        margin=ft.margin.only(bottom=30),
    )
    # Стиль для кнопок меню (выносим в переменную для повторного использования)
    menu_button_style = lambda color: ft.ButtonStyle(
        shape=ft.RoundedRectangleBorder(radius=8),
        bgcolor={
            ft.MaterialState.DEFAULT: "transparent",
            ft.MaterialState.HOVERED: color,
        },
        padding=10,
        overlay_color=ft.colors.TRANSPARENT,
    )

    def menu_button_style(color):
        return ft.ButtonStyle(
            shape=ft.RoundedRectangleBorder(radius=8),
            bgcolor={
                "": "transparent",  # Стандартное состояние
                "hovered": color,   # При наведении
            },
            padding=10,
        )

    # Создаем MenuBar с актуальными параметрами
    menubar = ft.MenuBar(
        expand=True,
        style=ft.MenuStyle(
            alignment=ft.alignment.top_left,
            bgcolor=COLORS["ACCENT"],
        ),
        controls=[
            # Меню "Информация"
            ft.SubmenuButton(
                content=ft.Text("Информация", color=COLORS["dark_purple"]),
                controls=[
                    ft.MenuItemButton(
                        content=ft.Text("О программе", color=COLORS["dark_purple"]),
                        leading=ft.Icon(ft.Icons.INFO, color=COLORS["dark_purple"]),
                        style=menu_button_style("#E2D4F0"),
                        on_click=lambda e: page.launch_url("https://example.com/about"),
                    ),
                    ft.MenuItemButton(
                        content=ft.Text("Выход", color=COLORS["dark_purple"]),
                        leading=ft.Icon(ft.Icons.EXIT_TO_APP, color=COLORS["dark_purple"]),
                        style=menu_button_style("#E2D4F0"),
                        on_click=lambda e: page.window_close,
                    ),
                ],
            ),

            # Основное меню расчетов
            ft.SubmenuButton(
                content=ft.Text("Расчёты", color=COLORS["dark_purple"]),
                controls=[
                    # Позиционирование
                    ft.SubmenuButton(
                        content=ft.Text("Позиционирование в пространстве", color=COLORS["dark_purple"]),
                        leading=ft.Image(
                            src="C:/Users/wertei siz/Dropbox/ПК/Desktop/ПО/Иконки флет/Позиционирование.png",
                            width=30,
                            height=30,
                        ),
                        controls=[
                            ft.MenuItemButton(
                                content=ft.Text("Определение координат", color=COLORS["dark_purple"]),
                                leading=ft.Image(
                                    src="C:/Users/wertei siz/Dropbox/ПК/Desktop/ПО/Иконки флет/подраздел.png",
                                    width=25,
                                    height=25,
                                ),
                                style=menu_button_style("#fcef72"),
                                on_click=positioning,
                            ),
                        ],
                    ),

                    # Энергетический расчет
                    ft.SubmenuButton(
                        content=ft.Text("Энергетический расчёт", color=COLORS["dark_purple"]),
                        leading=ft.Image(
                            src="C:/Users/wertei siz/Dropbox/ПК/Downloads/система.png",
                            width=30,
                            height=30,
                        ),
                        controls=[
                            ft.MenuItemButton(
                                content=ft.Text("Линия ЗС-ИЗС", color=COLORS["dark_purple"]),
                                leading=ft.Image(
                                    src="C:/Users/wertei siz/Dropbox/ПК/Desktop/ПО/Иконки флет/Земля.png",
                                    width=25,
                                    height=25,
                                ),
                                style=menu_button_style("#DA82FF"),
                                on_click=energy_calculation_up,
                            ),
                            ft.MenuItemButton(
                                content=ft.Text("Линия ИЗС-ЗС", color=COLORS["dark_purple"]),
                                leading=ft.Image(
                                    src="C:/Users/wertei siz/Dropbox/ПК/Downloads/спутник.png",
                                    width=25,
                                    height=25,
                                ),
                                style=menu_button_style("#F882FE"),
                                on_click=energy_calculation_down,
                            ),
                        ],
                    ),

                    # СППР
                    ft.SubmenuButton(
                        content=ft.Text("СППР", color=COLORS["dark_purple"]),
                        leading=ft.Image(
                            src="C:/Users/wertei siz/Dropbox/ПК/Desktop/ПО/Иконки флет/calculate.png",
                            width=30,
                            height=30,
                        ),
                        controls=[
                            ft.MenuItemButton(
                                content=ft.Text("Диаметр рефлектора", color=COLORS["dark_purple"]),
                                leading=ft.Image(
                                    src="C:/Users/wertei siz/Dropbox/ПК/Desktop/ПО/Иконки флет/циркуль.png",
                                    width=25,
                                    height=25,
                                ),
                                style=menu_button_style("#e583ac"),
                                on_click=antenna_diameter_decision,
                            ),
                        ],
                    ),
                ],
            ),
            # Меню "Файл"
            ft.SubmenuButton(
                content=ft.Text("Файл", color=COLORS["dark_purple"]),
                controls=[
                    ft.MenuItemButton(
                        content=ft.Text("Сохранить в TXT", color=COLORS["dark_purple"]),
                        leading=ft.Icon(ft.Icons.SAVE, color=COLORS["dark_purple"]),
                        style=menu_button_style("#E2F0D4"),
                        on_click=lambda e: save_result_txt(e),
                    ),
                ],
            ),
        ]
    )

    # Добавляем в страницу
    page.add(
        ft.Column(
            [menubar],
            expand=True,
        )
    )
ft.app(target=main)